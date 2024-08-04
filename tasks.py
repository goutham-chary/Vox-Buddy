from datetime import datetime
import random
import subprocess
import time
import pywhatkit
import wikipediaapi
import wikipedia
import speech_recognition as sr
import pyttsx3
import sys
import tkinter as tk
from tkinter import simpledialog
import subprocess
import webbrowser
import pyttsx3
import pyautogui
import requests
import os
import platform
import pyjokes
from docx import Document
from datetime import datetime
import re
import geocoder
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import ssl
import calendar
import speech_recognition as sr


def recognize_speech():
    recognizer = sr.Recognizer()
    with sr.Microphone() as source:
        print("Say something...")
        audio = recognizer.listen(source, 3)

    try:
        return recognizer.recognize_google(audio)
    except sr.UnknownValueError:
        print("Could not understand audio.")
    except sr.RequestError as e:
        print(f"Error connecting to Google API: {e}")
    return None

def open_ppt():
    try:
        note_name = ""
        if note_name:
            subprocess.run(["notepad.exe", note_name], check=True)
            speak(f"Opening note: {note_name}")
        else:
            speak("Note opening aborted. Could not understand the note name.")
    except subprocess.CalledProcessError as e:
        speak(f"Error opening note: {e}")
def greet():
    return "hello! How can I assist you today?"

def speak(text):
    engine = pyttsx3.init()
    engine.say(text)
    engine.runAndWait()

def tell_time():
    current_time = datetime.now().strftime("%H:%M")
    #speak(f"The current time is {current_time}")
    return f"The current time is {current_time}"

def get_weather(api_key, city):
    base_url = "http://api.openweathermap.org/data/2.5/weather"
    params = {
        'q': city,
        'appid': api_key,
        'units': 'metric',
    }

    response = requests.get(base_url, params=params)
    weather_data = response.json()

    if response.status_code == 200:
        main_weather = weather_data['weather'][0]['main']
        temperature = weather_data['main']['temp']
        description = weather_data['weather'][0]['description']

        return f"The weather in {city} is {main_weather}. Temperature: {temperature}°C. Description: {description}"
    else:
        return f"Error fetching weather information: {weather_data['message']}"

def manage_calendar(command):
    try:
        words = command.split()
        month = None
        year = None

        if "month" in words:
            month_index = words.index("month")
            if month_index + 1 < len(words):
                month = words[month_index + 1]

        if "year" in words:
            year_index = words.index("year")
            if year_index + 1 < len(words):
                year = words[year_index + 1]

        current_month, current_year = get_current_month_and_year()

        cal = calendar.monthcalendar(int(year) if year else current_year, int(month) if month else current_month)

        formatted_calendar = format_calendar(cal)

        return formatted_calendar
    except Exception as e:
        speak(f"Unable to manage calendar. Error: {str(e)}")
        return f"Unable to manage calendar. Error: {str(e)}"
def fetch_news():
    api_key = "8f2b367344084196b161ab5e0a05dba8"
    category='general'
    country='us'
    num_articles=10
    
    try:
        url = f'https://newsapi.org/v2/top-headlines?apiKey={api_key}&category={category}&country={country}&pageSize={num_articles}'
        response = requests.get(url)
        data = response.json()

        if response.status_code == 200:
            articles = data.get('articles', [])

            if not articles:
                speak("No news articles found.")
                return "No news articles found."

            news_output = "\n\n".join([f"{i + 1}. {article['title']} - {article['url']}" for i, article in enumerate(articles)])
            speak(f"Here are the latest news headlines:\n\n{news_output}")
            return f"Here are the latest news headlines:\n\n{news_output}"
        else:
            speak(f"Failed to fetch news. Error: {data.get('message', 'Unknown error')}")
            return f"Failed to fetch news. Error: {data.get('message', 'Unknown error')}"
    except Exception as e:
        speak(f"An error occurred: {str(e)}")
        return f"An error occurred: {str(e)}"
    
def format_calendar(cal):
    result = ""
    for week in cal:
        for day in week:
            if day == 0:
                result += "   "
            else:
                result += f"{day:2d} "
        result += "\n"
    return result

def get_current_month_and_year():
    import datetime
    current_date = datetime.datetime.now()
    return current_date.month, current_date.year

def extract_operands(command):
    words = command.split()
    try:
        index_of_and = words.index("and")
        operand1 = words[index_of_and - 1]
        operand2 = words[index_of_and + 1]
        return (operand1, operand2)
    except (ValueError, IndexError):
        return (None, None)
    
def tell_weather(api_key, city):
    weather_info = get_weather(api_key, city)
    #speak(weather_info)
    return weather_info

def tell_date_month_year_day(info_type):
    try:
        current_date = datetime.now()
        if info_type == "date":
            #speak(f"Today's date is {current_date.strftime('%Y-%m-%d')}")
            return f"Today's date is {current_date.strftime('%Y-%m-%d')}"
        elif info_type == "month":
            #speak(f"The current month is {current_date.strftime('%B')}")
            return f"The current month is {current_date.strftime('%B')}"
        elif info_type == "year":
            #speak(f"The current year is {current_date.strftime('%Y')}")
            return f"The current year is {current_date.strftime('%Y')}"
        elif info_type == "day":
            #speak(f"Today is {current_date.strftime('%A')}")
            return f"Today is {current_date.strftime('%A')}"
        else:
            speak("Invalid information type. Please specify date, month, year, or day.")
    except Exception as e:
        speak(f"Error telling the information: {e}")

def close_current_task():
    pyautogui.hotkey("alt", "f4")
    return("Closing the current task.")
    

def extract_range(command):
    words = command.split()
    try:
        index_of_between = words.index("between" or "and")
        start = words[index_of_between + 1]
        end = words[index_of_between + 3]
        return (start, end)
    except (ValueError, IndexError):
        return (None, None)
    
def generate_random_number(start, end):
    try:
        start = int(start)
        end = int(end)

        if start <= end:
            random_number = random.randint(start, end)
            speak(f"The random number between {start} and {end} is: {random_number}")
        else:
            speak("Invalid range. Please provide a valid range with the start value less than or equal to the end value.")
    except ValueError:
        speak("Invalid range. Please provide valid numeric values for the start and end of the range.")
    except Exception as e:
        speak(f"Error generating random number: {e}")

        
def send_whatsapp_message():
    try:
        app_name = "Whatsapp"
        speak(f"opening {app_name}")
        pyautogui.press('super')
        time.sleep(1)
        pyautogui.typewrite(app_name)
        time.sleep(1)
        pyautogui.press('enter')
        time.sleep(8)  

        speak("To whom would you like to send a WhatsApp message?")
        contact_name = recognize_speech()

        speak(f"What message would you like to send to {contact_name}?")
        message_content = recognize_speech()

        pyautogui.hotkey("ctrl", "f") 
        pyautogui.write(contact_name, interval=0.1)
        time.sleep(2)
        pyautogui.press("down")
        pyautogui.press("enter")

        time.sleep(2)
        pyautogui.write(message_content, interval=0.1)
        pyautogui.press("enter")

        speak(f"WhatsApp message sent to {contact_name}.")

    except Exception as e:
        speak("An error occurred while sending the WhatsApp message. Please try again.")
        print(f"An error occurred: {e}")

def get_recipient_email():
    root = tk.Tk()
    root.withdraw()
    recipient_email = simpledialog.askstring("Recipient's Email", "Please enter the recipient's email address:")
    root.destroy()
    return recipient_email

def validate_email(email):
    return "@" in email and "." in email.split("@")[1]

def send_email_via_gmail():
    try:
        speak("Please provide the recipient's email address.")
        recipient_email = get_recipient_email()

        if not validate_email(recipient_email):
            speak("Invalid email address. Please provide a valid email.")
            return "Invalid email address."

        speak("What should be the subject of the email?")
        subject = recognize_speech()

        speak("What should be the body of the email?")
        body = recognize_speech()

        sender_username = "YOUR_EMAIL_ID"
        sender_password = "YOUR_EMAIL_PASSWORD"
        smtp_server = "smtp.gmail.com"
        smtp_port = 465
        
        message = MIMEMultipart()
        message["From"] = sender_username
        message["To"] = recipient_email
        message["Subject"] = subject

        context = ssl.create_default_context()
        with smtplib.SMTP_SSL(smtp_server, smtp_port, context=context) as server:
            server.login(sender_username, sender_password)

            message.attach(MIMEText(body, "plain"))

            server.sendmail(sender_username, recipient_email, message.as_string())
        speak("Email sent successfully.")
        return "Email succesfully sent"
    except Exception as e:
        speak("An error occurred while sending the email. Please try again.")
        print(f"An error occurred: {e}")
        return f"An error occured : {e}"
  
def restart():
    try:
        subprocess.run(["shutdown", "/r", "/t", "0"], check=True)
        return "Restarting the laptop."
    except subprocess.CalledProcessError as e:
        return f"Error restarting the laptop: {e}"

def shutdown():
    try:
        subprocess.run(["shutdown", "/s", "/t", "0"], check=True)
        return "Shutting down the laptop."
    except subprocess.CalledProcessError as e:
        return f"Error shutting down the laptop: {e}"

def calculate_long_calculation(expression):
    try:
        expression = re.sub(r'\bplus\b', '+', expression, flags=re.IGNORECASE)
        expression = re.sub(r'\bminus\b', '-', expression, flags=re.IGNORECASE)
        expression = re.sub(r'\btimes\b', '*', expression, flags=re.IGNORECASE)
        expression = re.sub(r'\binto\b', '*', expression, flags=re.IGNORECASE)
        expression = re.sub(r'\bby\b', '/', expression, flags=re.IGNORECASE)

        print("Modified expression:", expression) 
        result = eval(expression)
        print(f"The result of the calculation is: {result}")
        return result
    except Exception as e:
        print(f"Error calculating: {e}")

def extract_calculation(command):
    try:
        index_of_calculate = command.lower().index("calculate") + len("calculate")
        calculation = command[index_of_calculate:].strip()
        return calculation
    except ValueError:
        return None
    
def lock_screen():
    system_platform = platform.system().lower()
    if system_platform == "windows":
        try:
            subprocess.run(["rundll32.exe", "user32.dll,LockWorkStation"], check=True)
            return "Locking the screen."
        except subprocess.CalledProcessError as e:
            return f"Error locking the screen: {e}"
    elif system_platform == "linux":
        try:
            subprocess.run(["gnome-screensaver-command", "--lock"], check=True)
            return "Locking the screen."
        except subprocess.CalledProcessError as e:
            return f"Error locking the screen: {e}"
    else:
        return "Lock screen not supported on this platform."
    
def open_social_media(platform):
    if platform.lower() == "instagram":
        url = "https://www.instagram.com/"
    elif platform.lower() == "twitter":
        url = "https://twitter.com/"
    elif platform.lower() == "facebook":
        url = "https://www.facebook.com/"
    else:
        return f"Unsupported social media platform: {platform}. Cannot open."

    webbrowser.open(url)
    return f"Opening {platform}."

def perform_arithmetic_operation(operation, operand1, operand2):
    try:
        operand1 = float(operand1)
        operand2 = float(operand2)

        if operation == "addition":
            result = operand1 + operand2
            speak(f"The result of addition is: {result}")
        elif operation == "subtraction":
            result = operand1 - operand2
            speak(f"The result of subtraction is: {result}")
        elif operation == "multiplication":
            result = operand1 * operand2
            speak(f"The result of multiplication is: {result}")
        elif operation == "division":
            if operand2 != 0:
                result = operand1 / operand2
                speak(f"The result of division is: {result}")
            else:
                speak("Cannot divide by zero. Please provide a non-zero divisor.")
        else:
            speak("Invalid arithmetic operation. Please specify addition, subtraction, multiplication, or division.")
    except ValueError:
        speak("Invalid operands. Please provide valid numeric operands.")
    except Exception as e:
        speak(f"Error performing arithmetic operation: {e}")

def extract_app_name(command):
    words = command.split()
    app_index = words.index("open") + 1
    if app_index < len(words):
        return words[app_index]
    else:
        return None

def open_application(app_name):
    speak(f"opening {app_name}")
    pyautogui.press('super')
    time.sleep(1)
    pyautogui.typewrite(app_name)
    time.sleep(1)
    pyautogui.press('enter')
    return f"opening {app_name}"


def open_chrome():
    speak("Opening Google Chrome.")
    webbrowser.open("www.google.com")

def open_google():
    speak("Opening Google")
    webbrowser.open("www.google.com")

def search_wikipedia(query):
    try:
        speak("What would you like to search on Wikipedia?")
        query = recognize_speech()

        wiki_wiki = wikipediaapi.Wikipedia('english')
        page_py = wiki_wiki.page(query)

        if not page_py.exists():
            return "Sorry, no information found on Wikipedia for '{}'.".format(query)

        summary = page_py.text.split('.')
        result = '.'.join(summary[:5]) + '.'

        speak(f"According to Wikipedia, {result}")
        return (f"According to Wikipedia, {result}")

    except wikipedia.DisambiguationError as e:
        speak("There are multiple results for your query. Please be more specific.")
        print(f"DisambiguationError: {e}")
        return (f"DisambiguationError: {e}")
    except wikipedia.PageError as e:
        speak("No results found on Wikipedia for your query.")
        print(f"PageError: {e}")
        return (f"PageError: {e}")
    except Exception as e:
        speak("An error occurred while trying to search on Wikipedia. Please try again.")
        print(f"An error occurred: {e}")
        return(f"An error occurred: {e}")
    

def play_song():
    speak("Please provide the name of the song you'd like to play.")
    song_name = recognize_speech()

    speak(f"Playing {song_name} on YouTube.")
    pywhatkit.playonyt(song_name)
    return speak(f"Playing {song_name} on YouTube.")

def control_volume(action):
    if action == "up":
        pyautogui.press('volumeup')
        return "Volume increased"
    elif action == "down":
        pyautogui.press('volumedown')
        return "Volume decreased"
    else:
        return "Invalid volume control command"
    
def search_product(platform, product):
    if platform.lower() == "amazon":
        url = f"https://www.amazon.in/s?k={'+'.join(product.split())}"
    elif platform.lower() == "flipkart":
        url = f"https://www.flipkart.com/search?q={'+'.join(product.split())}"
    else:
        return f"Unsupported platform for product search: {platform}. Cannot search."

    webbrowser.open(url)
    return f"Searching for {product} on {platform}."


def tell_joke():
    joke_to_tell = pyjokes.get_joke()
    #speak(joke_to_tell)
    return joke_to_tell

def open_mail():
    webbrowser.open("https://mail.google.com/")

def toggle_full_screen():
    pyautogui.press('f')
    return "Toggling full screen"

def write_note():
    try:
        speak("What is the name of the note you want to write in?")
        note_name = recognize_speech()
        if note_name:
            speak("What would you like to write in the note?")
            content = recognize_speech()
            with open(note_name, 'a') as note_file:
                note_file.write(content + '\n')
            speak(f"Content added to note: {note_name}")
        else:
            speak("Writing to note aborted. Could not understand the note name.")
    except Exception as e:
        speak(f"Error writing to note: {e}")

def create_note():
    try:
        speak("What would you like to name the new note?")
        note_name = recognize_speech()
        if note_name:
            note_path = f"{note_name}.txt"
            with open(note_path, 'w'):
                pass
            speak(f"Note created with the name: {note_name}.")
        else:
            speak("Note creation aborted. Could not understand the note name.")
    except subprocess.CalledProcessError as e:
        speak(f"Error creating note: {e}")

def open_note():
    try:
        speak("What is the name of the note you want to open?")
        note_name = recognize_speech()
        if note_name:
            subprocess.run(["notepad.exe", note_name], check=True)
            speak(f"Opening note: {note_name}")
        else:
            speak("Note opening aborted. Could not understand the note name.")
    except subprocess.CalledProcessError as e:
        speak(f"Error opening note: {e}")

def delete_note():
    try:
        speak("What is the name of the note you want to delete?")
        note_name = recognize_speech()
        if note_name:
            os.remove(f"{note_name}.txt")
            speak(f"Note deleted: {note_name}")
        else:
            speak("Note deletion aborted. Could not understand the note name.")
    except FileNotFoundError:
        speak(f"Note not found: {note_name}")
    except Exception as e:
        speak(f"Error deleting note: {e}")

def search_in_youtube(query):
    search_url = f"https://www.youtube.com/results?search_query={'+'.join(query.split())}"
    webbrowser.open(search_url)
    return f"Searching in YouTube: {query}"

def create_word_file():
    try:
        speak("What would you like to name the new Word file?")
        file_name = recognize_speech()
        if file_name:
            document = Document()
            document.save(file_name + ".docx")
            speak(f"Word file created with the name: {file_name}")
        else:
            speak("Word file creation aborted. Could not understand the file name.")
    except Exception as e:
        speak(f"Error creating Word file: {e}")

def open_word_file():
    try:
        speak("What is the name of the Word file you want to open?")
        file_name = recognize_speech()
        if file_name:
            os.system(f'start WINWORD "{file_name}.docx"')
            speak(f"Opening Word file: {file_name}")
        else:
            speak("Word file opening aborted. Could not understand the file name.")
    except subprocess.CalledProcessError as e:
        speak(f"Error opening Word file: {e}")

def get_current_location():
    try:
        location = geocoder.ip('me')

        city = location.city
        state = location.state
        country = location.country

        speak(f"You are currently in {city}, {state}, {country}.")
        return f"You are currently in {city}, {state}, {country}."
    except Exception as e:
        speak(f"Unable to retrieve current location. Error: {str(e)}")
        return f"Unable to retrieve current location. Error: {str(e)}"
    
def write_to_word_file():
    try:
        speak("What is the name of the Word file you want to write in?")
        file_name = recognize_speech()
        if file_name:
            speak("What would you like to write in the Word file?")
            content = recognize_speech()
            document = Document()
            document.add_paragraph(content)
            document.save(f"{file_name}.docx")
            speak(f"Content added to Word file: {file_name}")
        else:
                speak("Writing to Word file aborted. Could not understand the file name.")
    except Exception as e:
            speak(f"Error writing to Word file: {e}")

def delete_word_file():
    try:
        speak("What is the name of the Word file you want to delete?")
        file_name = recognize_speech()
        if file_name:
            os.remove(f"{file_name}.docx")
            speak(f"Word file deleted: {file_name}")
        else:
            speak("Word file deletion aborted. Could not understand the file name.")
    except FileNotFoundError:
        speak(f"Word file not found: {file_name}")
    except Exception as e:
        speak(f"Error deleting Word file: {e}")

def exit_voice_assistant():
    speak("Exiting Voice Assistant. Goodbye!")
    sys.exit()
